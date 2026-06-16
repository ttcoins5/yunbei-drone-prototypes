from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION, WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "奉飞飞后台已有功能模块菜单.docx"
SHOT_DIR = ROOT / "docs" / "screenshots" / "backend-modules"


MENU_MODULES = [
    {
        "menu": "工作台",
        "pages": ["工作台"],
        "fields": "今日新增订单、待派单订单、待审核飞手、待处理发票、指标明细、分页数据",
        "functions": "切换指标明细、跳转订单/飞手/发票处理页面、复制页面链接",
        "miniapp": "不直接搬到小程序；可转化为运营内部入口或飞手工作台统计。",
    },
    {
        "menu": "首页配置",
        "pages": ["首页配置"],
        "fields": "轮播素材、导航图片、标题、副标题、入口类型、跳转类型、跳转目标、启用状态、排序",
        "functions": "维护小程序首页轮播与 4 个大卡片、4 个小入口；手机预览同步更新",
        "miniapp": "直接决定小程序「首页」内容结构与入口展示。",
    },
    {
        "menu": "用户管理",
        "pages": ["用户列表", "用户详情"],
        "fields": "头像、昵称、手机号、性别、生日、地区、注册时间、订单记录、发票记录",
        "functions": "查询用户、查看个人信息、查看用户订单与发票记录",
        "miniapp": "映射到「我的 / 个人资料」，后台只做查询和客服支撑。",
    },
    {
        "menu": "商品管理",
        "pages": ["商品分类管理", "商品列表", "商品编辑"],
        "fields": "分类图标、分类名称、说明、商品数、启用状态、商品图、编号、名称、分类、规格、业务属性、价格、轮播图、介绍、评价",
        "functions": "分类增删改查、启停排序；商品新建/编辑/删除；维护多规格、预约、支付、飞手服务属性",
        "miniapp": "映射到小程序「首页服务入口 / 商品列表 / 商品详情 / 确认订单」。",
    },
    {
        "menu": "订单管理",
        "pages": ["订单列表", "订单详情"],
        "fields": "订单号、用户、商品/服务、金额、是否需要飞手、状态、预约日期、时段、联系电话、地址、备注、备注照片、已分配飞手",
        "functions": "筛选订单、查看详情、预览备注照片、分配/调整飞手、查看履约状态",
        "miniapp": "映射到底部导航「订单」及订单详情、继续支付、查看进度、评价/发票入口。",
    },
    {
        "menu": "培训管理",
        "pages": ["培训报名列表", "培训报名详情"],
        "fields": "线索编号、联系人、课程意向、手机号、报名时间、备注、来源",
        "functions": "查询培训线索、查看报名详情",
        "miniapp": "可放入「首页」培训入口或「我的」培训报名记录，首版可后置。",
    },
    {
        "menu": "飞手管理",
        "pages": ["入驻申请", "飞手审核详情", "已认证飞手", "飞手详情"],
        "fields": "申请编号、申请人、主体、公司、申请时间、状态、姓名、手机号、生日、区域、身份证、执照、无人机照片、机型、序列号、UAS 码、公司资料、服务区域、设备、关联订单",
        "functions": "筛选申请、审核通过/驳回、查看认证资料、查看飞手履约订单",
        "miniapp": "映射到「我的 / 飞手认证」和飞手工作台；普通用户认证通过后出现飞手入口。",
    },
    {
        "menu": "飞行报备管理",
        "pages": ["飞行报备列表", "飞行报备详情"],
        "fields": "报备编号、飞手、机型、架次、飞行时长、报备时间、状态、飞行区域、日期、时段、备注、来源、序列号、UAS 码",
        "functions": "筛选报备、查看详情、确认报备、推送第三方",
        "miniapp": "映射到飞手侧「飞行报备」填写、记录和详情。",
    },
    {
        "menu": "任务需求与意愿",
        "pages": ["任务需求列表", "发布任务", "任务详情", "意愿名单"],
        "fields": "需求编号、标题、服务时间、地址、状态、意愿人数、备注、富文本说明、飞手、主体、区域、设备、提交时间",
        "functions": "发布任务、查看任务详情、查看参与意愿名单",
        "miniapp": "映射到飞手侧「任务」列表、任务详情、提交/取消参与意愿。",
    },
    {
        "menu": "发票中心",
        "pages": ["发票列表", "发票详情"],
        "fields": "申请编号、用户、发票抬头、订单数、金额、申请时间、状态、发票类型、税号、邮箱、关联订单、驳回原因、发票文件",
        "functions": "查询发票申请、审核通过/驳回、上传发票文件",
        "miniapp": "映射到「我的 / 发票申请、发票记录、发票详情」。",
    },
    {
        "menu": "关于我们",
        "pages": ["关于我们"],
        "fields": "企业名称、Logo、企业介绍、电话、地址等富文本内容",
        "functions": "编辑企业介绍、预览、保存",
        "miniapp": "映射到「我的 / 关于我们」或首页品牌介绍入口。",
    },
]


SCREENSHOTS = [
    ("01-homepage-nav.png", "首页配置：轮播、导航入口和手机预览"),
    ("02-users.png", "用户列表：用户基础资料字段"),
    ("03-products.png", "商品列表：商品基础字段与筛选"),
    ("04-product-edit.png", "商品编辑：规格、业务属性、富文本介绍"),
    ("05-orders.png", "订单列表：订单筛选、状态和派单入口"),
    ("06-order-detail.png", "订单详情：预约资料、飞手分配和状态流转"),
    ("07-pilot-review.png", "飞手审核：个人、公司、资质和设备资料"),
    ("08-flight-reports.png", "飞行报备：报备列表和状态字段"),
    ("09-tasks.png", "任务需求：任务信息和意愿名单入口"),
    ("10-invoice-detail.png", "发票详情：票据信息、关联订单和审核动作"),
    ("11-about.png", "关于我们：企业介绍富文本配置"),
]


def set_run_font(run, name="Calibri", size=None, color=None, bold=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_widths(table, widths):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_paragraph(doc, text="", style=None, size=11, bold=False, color=None, after=6):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    if text:
        run = p.add_run(text)
        set_run_font(run, size=size, color=color, bold=bold)
    return p


def add_heading(doc, text, level):
    p = doc.add_heading("", level=level)
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(7 if level == 1 else 5)
    run = p.add_run(text)
    set_run_font(run, size=16 if level == 1 else 13, color="2E74B5" if level <= 2 else "1F4D78", bold=True)
    return p


def style_table_text(table, header_rows=1):
    for r_idx, row in enumerate(table.rows):
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.line_spacing = 1.15
                for run in p.runs:
                    set_run_font(run, size=9.2 if r_idx else 9.5, bold=(r_idx < header_rows))
            if r_idx < header_rows:
                set_cell_shading(cell, "E8EEF5")


def add_module_table(doc):
    table = doc.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    headers = ["左侧菜单", "页面名称", "主要字段", "后台功能", "小程序映射建议"]
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
    for item in MENU_MODULES:
        row = table.add_row().cells
        row[0].text = item["menu"]
        row[1].text = "\n".join(item["pages"])
        row[2].text = item["fields"]
        row[3].text = item["functions"]
        row[4].text = item["miniapp"]
    set_table_widths(table, [1300, 1450, 2700, 2200, 1710])
    style_table_text(table)


def add_screenshot_section(doc):
    first_section = doc.add_section(WD_SECTION.NEW_PAGE)
    first_section.orientation = WD_ORIENT.LANDSCAPE
    first_section.page_width = Inches(11)
    first_section.page_height = Inches(8.5)
    first_section.top_margin = Inches(0.45)
    first_section.right_margin = Inches(0.45)
    first_section.bottom_margin = Inches(0.45)
    first_section.left_margin = Inches(0.45)
    add_heading(doc, "字段截图索引", 1)
    add_paragraph(
        doc,
        "以下截图来自当前后台原型页面，用于辅助识别字段区、筛选区、表格和详情布局。字段以正文清单为准，截图作为视觉佐证。",
        size=10.5,
    )
    for idx, (filename, caption) in enumerate(SCREENSHOTS, start=1):
        image_path = SHOT_DIR / filename
        if not image_path.exists():
            continue
        section = doc.add_section(WD_SECTION.NEW_PAGE)
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Inches(11)
        section.page_height = Inches(8.5)
        section.top_margin = Inches(0.45)
        section.right_margin = Inches(0.45)
        section.bottom_margin = Inches(0.45)
        section.left_margin = Inches(0.45)
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(f"图 {idx}  {caption}")
        set_run_font(r, size=10.5, color="1F4D78", bold=True)
        doc.add_picture(str(image_path), width=Inches(9.75))


def build():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    header = section.header.paragraphs[0]
    header.text = "奉飞飞后台已有功能模块菜单"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        set_run_font(run, size=9, color="555555")

    footer = section.footer.paragraphs[0]
    footer.text = "用于小程序原型拆解"
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in footer.runs:
        set_run_font(run, size=9, color="555555")

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run("奉飞飞后台已有功能模块菜单")
    set_run_font(run, size=23, color="000000", bold=True)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    run = subtitle.add_run("面向微信小程序原型拆解的后台菜单、页面字段与截图索引")
    set_run_font(run, size=12, color="555555", bold=False)

    meta = doc.add_table(rows=3, cols=2)
    meta.style = "Table Grid"
    meta.rows[0].cells[0].text = "项目"
    meta.rows[0].cells[1].text = "奉飞飞无人机后台管理原型"
    meta.rows[1].cells[0].text = "整理日期"
    meta.rows[1].cells[1].text = "2026-06-15"
    meta.rows[2].cells[0].text = "整理用途"
    meta.rows[2].cells[1].text = "根据后台已有功能模块，辅助拆解微信小程序：首页、订单、消息、我的"
    set_table_widths(meta, [1700, 7660])
    style_table_text(meta)

    add_heading(doc, "使用说明", 1)
    add_paragraph(
        doc,
        "本清单不是把后台菜单原样搬到小程序，而是先识别后台已有业务能力，再筛选适合移动端用户、飞手或客户自助操作的功能。",
        size=10.5,
    )
    p = add_paragraph(doc, "建议小程序底部导航采用：首页、订单、消息、我的。商品、培训、飞手认证、报备、发票等作为二级页面从对应 Tab 进入。", size=10.5)
    p.runs[0].bold = True

    add_heading(doc, "左侧菜单与页面字段清单", 1)
    add_module_table(doc)

    add_heading(doc, "小程序四个 Tab 的初步承接关系", 1)
    tab_table = doc.add_table(rows=1, cols=3)
    tab_table.style = "Table Grid"
    for i, text in enumerate(["小程序导航", "承接后台能力", "建议页面"]):
        tab_table.rows[0].cells[i].text = text
    rows = [
        ("首页", "首页配置、商品分类、商品列表、关于我们、培训报名入口", "首页、服务分类、商品列表、商品详情、培训报名、关于我们"),
        ("订单", "订单管理、飞手分配状态、订单履约状态、发票关联订单", "订单列表、订单详情、确认订单、支付结果、评价、申请发票"),
        ("消息", "订单状态变更、飞手审核结果、任务通知、报备确认、发票审核结果", "消息列表、消息详情、分类筛选、已读/未读状态"),
        ("我的", "用户管理、飞手管理、发票中心、培训记录、关于我们", "个人资料、飞手认证、飞手工作台、发票记录、培训记录、设置"),
    ]
    for row_data in rows:
        row = tab_table.add_row().cells
        for i, text in enumerate(row_data):
            row[i].text = text
    set_table_widths(tab_table, [1500, 3650, 4210])
    style_table_text(tab_table)

    add_screenshot_section(doc)
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
