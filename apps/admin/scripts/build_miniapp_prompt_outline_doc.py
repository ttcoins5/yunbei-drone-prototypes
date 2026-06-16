from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION, WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "奉飞飞小程序原型产品提示词大纲.docx"

TAB_ROWS = [
    (
        "首页",
        "首页配置由后台维护；底部固定导航第 1 个入口。",
        "顶部轮播 / 品牌 Banner、后台配置的服务类商品分类入口、固定小模块：飞手加入、飞行报备、热销商品列表。",
        "轮播图、分类图标、分类名称、商品封面、商品名称、商品标签、起售价、销量、是否需要飞手服务、是否需要预约、是否需要在线支付。",
        "点击轮播或分类进入对应页面；点击飞手加入进入认证流程；点击飞行报备进入报备流程；点击热销商品进入商品详情。",
    ),
    (
        "订单",
        "订单页模块信息已提炼；底部固定导航第 2 个入口。",
        "顶部标题、订单状态 Tab、订单卡片列表、订单详情、后台派单与履约进度。",
        "订单号、商品图、商品名称、规格/服务类型、金额、状态标签、按钮：查看详情。状态包括：全部、待付款、待接单、待服务、待评价、已完成。",
        "Tab 切换筛选列表；点击查看详情进入订单详情；待付款可继续支付；待接单由后台管理员派单；待服务显示一个或多个飞手履约进度；待评价进入评价。",
    ),
    (
        "消息",
        "消息中心模块信息已提炼；底部固定导航第 3 个入口。",
        "消息卡片列表，首版只做列表，展示服务通知、优惠活动、在线客服、订单状态通知、审核结果通知、开票结果通知。",
        "消息类型图标/简称、消息标题、摘要、时间、未读红点、消息来源、关联业务编号。",
        "点击消息进入消息详情；未读消息点击后变已读；客服消息进入在线客服；订单、审核、开票、任务大厅消息跳转对应详情。",
    ),
    (
        "我的",
        "个人中心模块信息已提炼；底部固定导航第 4 个入口。",
        "用户头像昵称区、我的订单状态快捷入口、更多服务 8 个小模块。",
        "头像、昵称、个人资料入口；订单快捷状态：待付款、待接单、待服务、待评价、已完成；八个小模块：地址簿、联系客服、我的开票、意见反馈、关于我们、飞手加入、城市运营申请、任务大厅。",
        "点击头像进入个人资料；点击订单状态进入订单筛选；点击八个模块进入对应二级页面；飞手加入支持个人/企业主体；任务大厅仅认证飞手可见。",
    ),
]

BUSINESS_FLOW_ROWS = [
    (
        "阶段一：商品配置与飞手入驻",
        "后台管理员、飞手",
        "后台创建服务类商品分类，配置面向用户类型、是否需要飞手服务、是否需要预约、是否需要在线支付；创建或编辑商品并维护名称、详情、价格、上下架。飞手提交入驻申请，区分个人飞手和企业飞手。",
        "飞手字段：申请人、联系电话、出生年月、所在区域、身份证正反面、操作执照、无人机照片、机型、序列号、唯一识别码。企业主体额外补充公司名称、联系电话、所在区域。",
        "后台审核通过后进入可分配飞手池；审核不通过退回修改资料后重新提交。",
    ),
    (
        "阶段二：用户下单",
        "小程序用户、后台管理员",
        "用户浏览分类与商品详情，选择商品或预约服务。若商品需要预约，填写服务时间、地址、手机号、备注照片；若支持在线支付，订单进入待付款并在支付成功后继续流转。",
        "订单判断链路：是否需要预约、是否支持在线支付、是否需要飞手服务。",
        "在线支付成功后触发服务通知；无需在线支付的订单直接生成并进入后续状态。",
    ),
    (
        "阶段三：接单与履约",
        "后台管理员、飞手、小程序用户",
        "需要飞手服务的订单进入待接单，由后台管理员在商品订单中分配一个或多个审核通过飞手；分配后进入待服务。无需飞手的订单由后台确认服务或商品交付完成。",
        "飞手端查看被分配订单，执行服务并提交完成；每位飞手完成后更新个人履约结果。",
        "全部已分配飞手完成后，订单进入待评价；任务大厅为独立需求收集，不关联商品订单。",
    ),
    (
        "阶段四：评价与完成",
        "小程序用户、后台管理员",
        "用户收到待评价通知后可提交评价；若超过评价期限，可按超时规则自动完成。",
        "评价字段建议：星级、评价内容、服务商品、关联订单、提交时间。",
        "用户已评价或超时后，订单进入已完成。",
    ),
    (
        "阶段五：合并开票",
        "小程序用户、后台管理员",
        "用户进入发票中心，勾选一个或多个已完成且未开票订单，提交开票申请。",
        "合并开票条件：同一用户、同一发票抬头、同一税率。一张发票可关联多个订单。",
        "校验通过后生成合并申请并审核开票；校验不通过返回重新选择订单；用户可查看或下载发票。",
    ),
]


PROMPT_SECTIONS = [
    (
        "一、项目背景",
        [
            "项目名称：奉飞飞无人机服务微信小程序原型。",
            "三角色业务视角：小程序用户、后台管理员、飞手。",
            "小程序面向：普通用户、潜在飞手、城市运营申请人；后台用于配置首页、商品、订单、飞手、报备、发票与内容。",
            "本次目标：基于后台已有功能模块，设计一个可点击的小程序原型，底部导航固定为：首页、订单、消息、我的。",
            "请不要照搬后台菜单，要按小程序用户视角重组功能。",
        ],
    ),
    (
        "二、首页大纲",
        [
            "首页内容由后台「首页配置」维护，包括轮播素材、服务分类入口、跳转目标和启用状态。",
            "首页从上到下结构：轮播/品牌 Banner、服务分类入口、固定两个小模块、热销商品。",
            "固定两个小模块：飞手加入、飞行报备；这两个模块不受后台普通导航入口排序影响，始终展示在首页中部。",
            "后台商品分类支持服务类商品，商品配置属性包括：面向用户类型、是否需要飞手服务、是否需要预约、是否需要在线支付。",
            "热销商品按销量排序自动生成，字段包括商品图、名称、价格、标签、服务属性和详情入口。",
            "需要补充：首页分类数量、是否展示搜索、是否展示城市定位。",
        ],
    ),
    (
        "三、订单大纲",
        [
            "订单页页面标题为「我的订单」，顶部为状态 Tab。",
            "Tab 包括：全部、待付款、待接单、待服务、待评价；我的页面快捷入口额外包含已完成。",
            "订单卡片字段：订单号、商品/服务图、商品/服务名称、规格或服务类型、金额、状态标签、查看详情按钮。",
            "订单判断链路：是否需要预约、是否支持在线支付、是否需要飞手服务。",
            "订单详情需要展示：订单状态流程、商品快照、预约日期、预约时段、联系人、电话、地址、备注、备注照片、已分配飞手、支付信息。",
            "待接单由后台管理员派单，可分配一个或多个审核通过飞手；无需飞手订单由后台确认服务或商品交付完成。",
            "需要补充：用户是否可取消订单、是否可退款/售后。",
        ],
    ),
    (
        "四、消息中心大纲",
        [
            "消息页页面标题为「消息中心」，以卡片列表呈现。",
            "首版消息类型：服务通知、优惠活动、在线客服。",
            "消息中心首版只做列表，不做分类 Tab。",
            "可扩展消息类型：订单通知、飞手审核通知、报备进度通知、发票审核通知、任务大厅通知、系统公告。",
            "关键触发消息：支付成功、待接单、待服务、待评价、飞手审核结果、开票结果、任务大厅通知。",
            "消息卡片字段：类型简称/图标、标题、摘要、时间、未读红点。",
            "需要补充：是否支持删除、在线客服是否接微信客服组件。",
        ],
    ),
    (
        "五、我的大纲",
        [
            "我的页页面标题为「个人中心」。",
            "顶部用户区：头像、昵称、点击查看个人资料。",
            "我的订单卡片：待付款、待接单、待服务、待评价、已完成，点击进入订单页并带状态筛选。",
            "更多服务 8 个小模块：地址簿、联系客服、我的开票、意见反馈、关于我们、飞手加入、城市运营申请、任务大厅。",
            "飞手加入支持个人飞手和企业飞手两种主体；企业飞手字段包括公司名称、联系电话、所在区域。",
            "城市运营申请字段：机构名称、申请人、联系电话、申请区域、身份证正反面、营业执照、协议勾选、提交申请；当前只做信息收集，不需要后台审核。",
            "任务大厅仅认证飞手可见。",
            "需要补充：是否显示手机号、飞手加入入口是否根据认证状态变化。",
        ],
    ),
    (
        "六、发票中心大纲",
        [
            "发票中心首版开放申请发票，从「我的 / 我的开票」进入。",
            "用户可勾选一个或多个已完成且未开票订单。",
            "合并开票条件：同一用户、同一发票抬头、同一税率。",
            "校验通过后生成合并申请并进入审核开票；校验不通过返回重新选择订单。",
            "用户可查看或下载发票。",
        ],
    ),
]


def set_run_font(run, name="Calibri", size=None, color=None, bold=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_widths(table, widths):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell_margins(cell)


def style_table(table, header_rows=1, font_size=9.2):
    for r_idx, row in enumerate(table.rows):
        for cell in row.cells:
            if r_idx < header_rows:
                shade(cell, "E8EEF5")
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.line_spacing = 1.15
                for run in p.runs:
                    set_run_font(run, size=font_size, bold=(r_idx < header_rows))


def add_heading(doc, text, level=1):
    p = doc.add_heading("", level=level)
    p.paragraph_format.space_before = Pt(18 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(8 if level == 1 else 5)
    r = p.add_run(text)
    set_run_font(r, size=16 if level == 1 else 13, color="2E74B5" if level == 1 else "1F4D78", bold=True)
    return p


def add_para(doc, text="", size=11, bold=False, color=None, after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    if text:
        r = p.add_run(text)
        set_run_font(r, size=size, color=color, bold=bold)
    return p


def add_real_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    r = p.add_run(text)
    set_run_font(r, size=10.5)
    return p


def build_tab_table(doc):
    table = doc.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    for i, header in enumerate(["导航", "定位", "页面结构", "主要字段", "关键交互"]):
        table.rows[0].cells[i].text = header
    for row_data in TAB_ROWS:
        row = table.add_row().cells
        for i, value in enumerate(row_data):
            row[i].text = value
    set_table_widths(table, [800, 1700, 2350, 2500, 2010])
    style_table(table, font_size=8.8)


def build_business_flow_table(doc):
    table = doc.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    for i, header in enumerate(["阶段", "参与角色", "流程说明", "关键字段 / 判断", "状态 / 结果"]):
        table.rows[0].cells[i].text = header
    for row_data in BUSINESS_FLOW_ROWS:
        row = table.add_row().cells
        for i, value in enumerate(row_data):
            row[i].text = value
    set_table_widths(table, [1450, 1250, 3100, 2050, 1510])
    style_table(table, font_size=8.4)


def add_prompt_block(doc):
    add_heading(doc, "可直接复制的产品提示词", 1)
    prompt = (
        "请基于奉飞飞后台已有功能模块，设计一个微信小程序高保真交互原型。"
        "业务需覆盖小程序用户、后台管理员、飞手三个角色，以及商品配置与飞手入驻、用户下单、接单与履约、评价与完成、合并开票五个阶段。"
        "小程序底部导航固定为：首页、订单、消息、我的。首页内容由后台配置，首页中部固定展示两个小模块：飞手加入、飞行报备；再往下展示按销量排序的热销商品。"
        "订单页、消息页、我的页的模块信息已提炼为文字说明。"
        "订单需体现是否预约、是否在线支付、是否需要飞手服务的判断链路；待接单由后台管理员派单，可分配一个或多个审核通过飞手。"
        "消息中心首版只做列表，包含支付成功、待接单、待服务、待评价、飞手审核结果、开票结果、任务大厅通知等触发消息。"
        "我的页需包含飞手加入、城市运营申请、任务大厅和发票中心；任务大厅仅认证飞手可见，发票中心首版开放申请发票并支持合并开票。"
        "请输出可点击页面结构、字段、状态、跳转关系、空状态、异常状态，并保持微信小程序业务工具型风格。"
    )
    p = add_para(doc, prompt, size=10.5)
    p.paragraph_format.left_indent = Inches(0.15)
    p.paragraph_format.right_indent = Inches(0.15)


def add_reference_images(doc):
    add_heading(doc, "已提炼模块信息", 1)
    add_para(doc, "以下为订单、消息、我的三个页面的模块信息整理，可直接放入产品提示词。", size=10.5)
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    for i, header in enumerate(["页面", "模块信息", "应输出到原型的文字要求"]):
        table.rows[0].cells[i].text = header
    rows = [
        (
            "订单",
            "顶部标题：我的订单；状态 Tab：全部、待付款、待接单、待服务、待评价；订单卡片包含订单号、状态标签、商品图、商品名称、规格/服务类型、价格、查看详情按钮；底部导航：首页、订单、消息、我的。",
            "订单页按状态 Tab 筛选；订单卡片统一展示订单号、商品图、商品名称、规格或服务类型、金额、状态、查看详情；状态色需区分待接单、待服务、已完成等。",
        ),
        (
            "消息",
            "顶部标题：消息中心；消息卡片：服务通知、优惠活动、在线客服；每条消息包含类型简称、标题、摘要、时间；服务通知有未读红点。",
            "消息中心首版至少包含服务通知、优惠活动、在线客服三类；支持未读红点、时间显示、点击进入详情或客服。",
        ),
        (
            "我的",
            "顶部标题：个人中心；用户区：头像、昵称、查看个人资料；我的订单：待付款、待接单、待服务、待评价、已完成；更多服务八个模块：地址簿、联系客服、我的开票、意见反馈、关于我们、飞手加入、城市运营申请、任务大厅。",
            "我的页需要用户资料入口、订单状态快捷入口和 8 个服务模块；点击订单状态进入订单页对应筛选，点击服务模块进入对应二级页面。",
        ),
    ]
    for row_data in rows:
        row = table.add_row().cells
        for i, value in enumerate(row_data):
            row[i].text = value
    set_table_widths(table, [900, 4230, 4230])
    style_table(table, header_rows=1, font_size=9)


def build():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    header = section.header.paragraphs[0]
    header.text = "奉飞飞小程序原型产品提示词大纲"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        set_run_font(run, size=9, color="555555")

    footer = section.footer.paragraphs[0]
    footer.text = "给原型设计师 / 给 Codex 使用"
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in footer.runs:
        set_run_font(run, size=9, color="555555")

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run("奉飞飞小程序原型产品提示词大纲")
    set_run_font(run, size=23, color="000000", bold=True)
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    run = subtitle.add_run("按首页、订单、消息、我的四个底部导航补充，可继续填写业务规则")
    set_run_font(run, size=12, color="555555")

    meta = doc.add_table(rows=3, cols=2)
    meta.style = "Table Grid"
    rows = [
        ("底部导航", "首页 / 订单 / 消息 / 我的"),
        ("信息来源", "后台首页配置、订单管理、订单/消息/我的页面模块信息"),
        ("整理日期", "2026-06-15"),
    ]
    for r, (label, value) in enumerate(rows):
        meta.rows[r].cells[0].text = label
        meta.rows[r].cells[1].text = value
    set_table_widths(meta, [1700, 7660])
    style_table(meta)

    add_heading(doc, "总目标", 1)
    for text in [
        "产出一个微信小程序原型，面向小程序用户、后台管理员、飞手三角色业务协作，同时在移动端呈现用户与飞手高频入口。",
        "后台负责配置首页轮播、导航入口、商品、订单、飞手、报备、发票和企业介绍；小程序只呈现移动端高频功能。",
        "业务流程覆盖商品配置与飞手入驻、用户下单、接单与履约、评价与完成、合并开票五个阶段。",
        "视觉风格：浅灰背景、白色圆角卡片、清晰状态标签、底部四栏导航。",
    ]:
        add_real_bullet(doc, text)

    add_heading(doc, "业务主流程", 1)
    build_business_flow_table(doc)

    add_heading(doc, "四个底部导航大纲", 1)
    build_tab_table(doc)

    add_prompt_block(doc)

    add_heading(doc, "分页面补充大纲", 1)
    for title_text, bullets in PROMPT_SECTIONS:
        add_heading(doc, title_text, 2)
        for bullet in bullets:
            add_real_bullet(doc, bullet)

    add_heading(doc, "已确认规则", 1)
    confirmed = [
        "首页热销商品按销量排序自动生成。",
        "订单待接单由后台管理员派单。",
        "消息中心首版只做列表，不做分类 Tab。",
        "飞手加入包含个人飞手和企业飞手两种主体。",
        "城市运营申请字段包括：机构名称、申请人、联系电话、申请区域、身份证正反面、营业执照、协议勾选、提交申请；当前只做信息收集，不需要后台审核。",
        "任务大厅仅认证飞手可见，且只收集飞手接单意愿，不关联商品订单。",
        "发票中心首版开放申请发票，支持一个或多个已完成且未开票订单合并开票。",
    ]
    for item in confirmed:
        add_real_bullet(doc, item)

    add_heading(doc, "剩余待补充信息", 1)
    questions = [
        "用户是否可取消订单、申请退款或售后。",
        "评价是否需要照片、匿名、审核或超时自动完成的具体天数。",
        "在线客服是否接入微信客服组件。",
        "首页是否展示搜索、城市定位、分类数量上限。",
        "飞手完成服务时是否需要上传成果照片、附件或定位信息。",
    ]
    for q in questions:
        add_real_bullet(doc, q)

    add_reference_images(doc)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
